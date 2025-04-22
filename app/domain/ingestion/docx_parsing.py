# app/utils/docx_parsing.py
"""
This module contains functions for parsing files.
"""
import io
from typing import List, Dict, Any

import pandas as pd
from fastapi import UploadFile, File
from langchain_core.documents.base import Document as LangchainDocument
import re
from docx import Document
from bs4 import BeautifulSoup

from api.logging_theme import setup_logger


class DocxParser:
    """
    Define Docx Parser
    """

    def __init__(self):
        self.results: List[LangchainDocument] = []
        self.logger = setup_logger(__name__)

    def transform_to_json_structure(self, tables: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """
        Transform tables data into a structured JSON format with proper header processing.

        This function processes tables extracted from Word documents, handling multi-line headers
        and removing duplications in header text. It converts tabular data into a JSON structure
        where each table includes its label and data rows.

        Args:
            tables: A list of dictionaries, each containing table data with rows and a label

        Returns:
            A dictionary with a "tables" key containing the processed table data
        """
        try:
            json_result = {"tables": []}

            for table in tables:
                rows = table["rows"]
                if not rows:
                    continue

                # Detect multi-line headers
                header_rows = []
                for i, row in enumerate(rows):
                    # Check if a row is a potential header (all cells filled and no numeric-only columns)
                    if all(row) and not any(re.search(r'^\d+$', cell) for cell in row):
                        header_rows.append(i)
                    else:
                        break

                # Process headers if there are multiple rows
                table_dict = []
                if header_rows:
                    # Merge multiple header rows WITHOUT duplication
                    headers = rows[header_rows[0]:header_rows[-1] + 1]

                    # Process each column to avoid duplication
                    merged_headers = []
                    for col in zip(*headers):
                        # Remove empty strings and duplicates while preserving order
                        unique_parts = []
                        for part in col:
                            part = part.strip()
                            if part and part not in unique_parts:
                                unique_parts.append(part)
                        merged_headers.append(" ".join(unique_parts))

                    # Process data rows starting after the headers
                    data_rows = rows[len(header_rows):]

                    # Only create dictionaries if we have both headers and data
                    if merged_headers and data_rows:
                        table_dict = [
                            {merged_headers[i]: row[i] for i in range(min(len(merged_headers), len(row)))}
                            for row in data_rows
                        ]
                else:
                    # Handle tables without headers
                    table_dict = rows

                # Add the processed table to the JSON result only if we have data
                if table_dict:  # Only add non-empty tables
                    json_result["tables"].append({
                        "label": table["label"],  # Include the label detected before the table
                        "data": table_dict
                    })

            return json_result
        except Exception as e:
            self.logger.error(f"Error extracting tables from Word document: {str(e)}")
            raise ValueError(f"Error extracting tables from Word document: {str(e)}")

    def extract_text_before_table(self, doc: Document, table_idx: int) -> str:
        """
        Extract the paragraph text that appears immediately before a table in a Word document.

        This function navigates through the document's structure to find the paragraph
        that appears directly before the specified table. This text is often a caption or
        description of the table.

        Args:
            doc: A python-docx Document object
            table_idx: The index of the target table in the document

        Returns:
            The text of the paragraph immediately preceding the table, or a default
            string like "Table X" if no such paragraph is found
        """
        try:
            # Get the target table
            table = doc.tables[table_idx]

            # Get all content elements in document order
            all_elements = []
            for element in doc._element.body:
                if element.tag.endswith('tbl'):  # Table element
                    all_elements.append(('table', element))
                elif element.tag.endswith('p'):  # Paragraph element
                    all_elements.append(('paragraph', element))

            # Find the position of our target table
            table_positions = [i for i, (type_el, el) in enumerate(all_elements)
                               if type_el == 'table']
            if table_idx >= len(table_positions):
                return f"Table {table_idx + 1}"

            target_table_pos = table_positions[table_idx]

            # Look backward to find the nearest paragraph
            prev_paragraphs = []
            for i in range(target_table_pos - 1, -1, -1):
                if i < 0:
                    break

                if all_elements[i][0] == 'paragraph':
                    # Convert XML element to paragraph and get text
                    for p in doc.paragraphs:
                        if p._element is all_elements[i][1]:
                            text = p.text.strip()
                            if text:  # If non-empty paragraph
                                return text

            # If no appropriate paragraph found
            return f"Table {table_idx + 1}"
        except Exception as e:
            self.logger.error(f"Error extracting tables from Word document: {str(e)}")
            raise ValueError(f"Error extracting tables from Word document: {str(e)}")

    def extract_tables_with_labels(self, doc_file: UploadFile = File(...)) -> List[Dict[str, Any]]:
        """
        Extract tables from a Word document along with their preceding labels.

        This function processes a Word document, extracting all tables and identifying
        the text that appears before each table as a potential label or caption.
        It normalizes the table structure by ensuring all rows have the same number of columns.

        Args:
            doc_file: the Word document file

        Returns:
            A list of dictionaries, each containing a table's label and row data
        """
        try:
            doc_content = io.BytesIO(doc_file.file.read())
            doc = Document(doc_content)
            extracted_data = []

            for table_idx, table in enumerate(doc.tables):
                table_label = self.extract_text_before_table(doc, table_idx)  # Get text before the table
                table_data = {"label": table_label, "rows": []}

                max_cols = max(len(row.cells) for row in table.rows) if table.rows else 0

                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    row_data += [""] * (max_cols - len(row_data))  # Pad columns if needed
                    table_data["rows"].append(row_data)

                extracted_data.append(table_data)

            return extracted_data
        except Exception as e:
            self.logger.error(f"Error extracting tables from Word document: {str(e)}")
            raise ValueError(f"Error extracting tables from Word document: {str(e)}")

    def convert_table_to_tsv_string(self, html_content: str) -> str:
        """
        Converts an HTML table to a TSV (Tab-Separated Values) string.
        
        Args:
            html_content (str): The HTML content containing the table.
        
        Returns:
            str: A TSV-formatted string representation of the table, 
                 or an error message if no table is found.
        """
        try:
            # Parse HTML content
            soup = BeautifulSoup(html_content, 'html.parser')

            # Find the table
            table = soup.find('table')

            if not table:
                self.logger.error("No table found in the HTML content.")
                raise ValueError("No table found in the HTML content.")

            # Extract all rows
            rows = table.find_all('tr')

            # Initialize result string
            result = ""

            # Process each row
            for row in rows:
                # Get all cells in this row
                cells = row.find_all(['td', 'th'])

                # Extract text from each cell
                row_texts = [cell.get_text().strip() for cell in cells]

                # Join cell texts with tabs
                row_string = "\t".join(row_texts)

                # Add row to result
                result += row_string + "\n"

            self.logger.info("Successfully converted HTML table to TSV format.")
            return result
        except Exception as e:
            self.logger.error(f"Error converting table to TSV: {str(e)}")
            raise ValueError(f"Error converting table to TSV: {str(e)}")
        
    async def faq_parsing(self, upload_file: UploadFile = File(...), question_column_name: str = "CÂU HỎI",
                          answer_column_name: str = "CÂU TRẢ LỜI ", header: int = 1) -> List[LangchainDocument]:
        """
        Reads a CSV file containing FAQ data and returns a list of LangchainDocument objects.

        Args:
            upload_file (UploadFile): The list of uploaded CSV files containing FAQ data.
            question_column_name (str): The column name for questions in the CSV file.
            answer_column_name (str): The column name for answers in the CSV file.

        Returns:
            List[LangchainDocument]: A list of LangchainDocument objects with page content set to the question
                                     and metadata containing the answer.
        """
        try:
            content = await upload_file.read()
            faq = pd.read_csv(io.BytesIO(content), header=header)
            self.logger.debug(faq)
            for item in faq.to_dict(orient="records"):
                self.results.append(LangchainDocument(page_content=item[question_column_name],
                                                      metadata={"answer": item[answer_column_name]}))
            self.logger.info(f"Successfully parsed {len(self.results)} FAQ data")
            return self.results
        except Exception as e:
            self.logger.error(f"Error parsing FAQ data: {str(e)}")
            raise ValueError(f"Error parsing FAQ data: {str(e)}")
